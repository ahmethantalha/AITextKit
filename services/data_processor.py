import os
import json
import time
from datetime import datetime
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import docx
from fpdf import FPDF

class DataProcessor:
    def __init__(self):
        self.supported_formats = {
            'pdf': ['.pdf'],
            'image': ['.jpg', '.jpeg', '.png'],
            'text': ['.txt'],
            'word': ['.docx'],
            'json': ['.json']
        }
        
        self.prompt_templates = {
            "Soru-Cevap Üretimi": """Aşağıdaki metni analiz et, {KONU} (fine tunning metoduyla) eğiteceğimiz yapay zekamızı en iyi şekilde eğitebileceğim şekilde, **özetlenmiş** olarak, **tekrarsız** bir şekilde **farklı** soru-cevap çiftleri üret.
**Önemli Not:** Aynı soru-cevap çiftlerini tekrar etme!

JSON formatında çıktı üret:
{
"soru-cevaplar": [
    {
    "soru": "örnek soru",
    "cevap": "örnek cevap"
    }
]
}""",
            "Metin Özeti Oluştur": """Aşağıdaki metni analiz et ve kapsamlı bir özet oluştur. 
Önemli noktaları, ana fikirleri ve temel argümanları içermelidir.
Paragraflar halinde, akıcı ve anlaşılır bir dilde yaz.""",
            "Özel Prompt": ""
        }
    
    def get_file_type(self, extension):
        """Determine file type from extension"""
        for file_type, extensions in self.supported_formats.items():
            if extension in extensions:
                return file_type
        return None
    
    def get_prompt(self, prompt_type, topic="", custom_prompt=""):
        """Get the appropriate prompt based on type and parameters"""
        if prompt_type == "Özel Prompt":
            # Özel prompt boşsa, basit bir varsayılan değer döndür
            return custom_prompt or "Lütfen metni analiz et ve yanıt ver."
        elif prompt_type == "Soru-Cevap Üretimi":
            return self.prompt_templates[prompt_type].replace("{KONU}", topic or "[KONU]")
        else:
            return self.prompt_templates[prompt_type]
    
    def process_file(self, file_path, file_type, use_vision=False):
        """Process a file based on its type and return text chunks"""
        if file_type == 'pdf':
            return self.process_pdf(file_path)
        elif file_type == 'image':
            # Vision API kullanılacaksa, resim dosya yolunu doğrudan döndür
            if use_vision:
                return [file_path]  # Vision API için resim yolunu döndür
            # Normal işlem için OCR kullan
            return self.process_image(file_path)
        elif file_type == 'text':
            return self.process_text(file_path)
        elif file_type == 'word':
            return self.process_word(file_path)
        elif file_type == 'json':
            return self.process_json(file_path)
        return []
    
    def process_pdf(self, file_path, chunk_size=5):
        """Extract text from PDF files in chunks"""
        try:
            doc = fitz.open(file_path)
            texts = []
            
            for i in range(0, len(doc), chunk_size):
                chunk = ""
                end_idx = min(i + chunk_size, len(doc))
                
                for j in range(i, end_idx):
                    chunk += doc[j].get_text()
                
                if chunk.strip():
                    texts.append(chunk.strip())
                    
            return texts
        except Exception as e:
            raise Exception(f"PDF işleme hatası: Dosya okunamadı veya hasar görmüş olabilir. Detay: {e}")

    def process_image(self, file_path):
        """Extract text from images using OCR"""
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang='tur')  # Using Turkish language for OCR
            return [text] if text.strip() else []
        except Exception as e:
            raise Exception(f"Görüntü işleme hatası: Dosya formatı desteklenmiyor veya hasar görmüş olabilir. Detay: {e}")

    def process_text(self, file_path):
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return [f.read()]
        except UnicodeDecodeError:
            # Try different encodings if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return [f.read()]
            except Exception as e2:
                raise Exception(f"Metin dosyası okuma hatası: Dosya encoding hatası. Detay: {e2}")
        except Exception as e:
            raise Exception(f"Metin dosyası okuma hatası: Dosya erişim sorunu olabilir. Detay: {e}")

    def process_word(self, file_path):
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            return ['\n'.join([paragraph.text for paragraph in doc.paragraphs])]
        except Exception as e:
            raise Exception(f"Word dosyası işleme hatası: Dosya formatı uyumsuz veya hasar görmüş olabilir. Detay: {e}")

    def process_json(self, file_path):
        """Extract content from JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return [json.dumps(data, ensure_ascii=False)]
        except UnicodeDecodeError:
            # Try different encodings if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    data = json.load(f)
                    return [json.dumps(data, ensure_ascii=False)]
            except Exception as e2:
                raise Exception(f"JSON dosyası okuma hatası: JSON formatı geçersiz. Detay: {e2}")
        except Exception as e:
            raise Exception(f"JSON dosyası okuma hatası: JSON formatı geçersiz veya encoding hatası olabilir. Detay: {e}")
    
    def save_content(self, content, output_format, save_dir):
        """Save content in the specified format"""
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        
        if output_format == "TXT":
            filename = f"ozet_{timestamp}.txt"
            file_path = os.path.join(save_dir, filename)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
                
        elif output_format == "DOCX":
            filename = f"ozet_{timestamp}.docx"
            file_path = os.path.join(save_dir, filename)
            
            doc = docx.Document()
            doc.add_heading('Metin Özeti', 0)
            for paragraph in content.split('\n\n'):
                doc.add_paragraph(paragraph)
            doc.save(file_path)
                
        elif output_format == "PDF":
            filename = f"ozet_{timestamp}.pdf"
            file_path = os.path.join(save_dir, filename)
            
            pdf = FPDF()
            pdf.add_page()
            
            # Try to use a font that supports Turkish characters
            try:
                pdf.add_font('DejaVu', '', 'DejaVuSansCondensed.ttf', uni=True)
                pdf.set_font('DejaVu', '', 12)
            except:
                # Default to Arial if DejaVu font is not available
                pdf.set_font("Arial", size=12)
            
            # Handle line breaks and wrap text
            for line in content.split('\n'):
                pdf.multi_cell(0, 10, txt=line)
            
            pdf.output(file_path)
        
        return filename
    
    def save_json(self, data, save_dir):
        """Save JSON data to a file"""
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        filename = f"soru_cevap_{timestamp}.json"
        file_path = os.path.join(save_dir, filename)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
            
        return filename
